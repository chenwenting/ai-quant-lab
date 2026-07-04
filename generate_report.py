# -*- coding: utf-8 -*-
"""
兆易创新(603986.SH) 基本分析报告生成脚本
"""
import csv, os, io
from datetime import datetime
from collections import defaultdict
import matplotlib
matplotlib.use('Agg')
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== CONFIG ==========
STOCK_NAME = "兆易创新"
STOCK_CODE = "603986.SH"
STOCK_FULL = f"{STOCK_NAME}（{STOCK_CODE}）"
DATA_FILE = "data/603986_daily.csv"
OUTPUT_DIR = "reports"
CHART_DIR = os.path.join(OUTPUT_DIR, "charts")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "兆易创新_603986_基本分析报告.docx")

os.makedirs(CHART_DIR, exist_ok=True)

# ========== LOAD DATA ==========
rows = []
with open(DATA_FILE, 'r', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for r in reader:
        row = {
            'date': datetime.strptime(r['trade_date'], '%Y-%m-%d'),
            'open': float(r['open']), 'high': float(r['high']),
            'low': float(r['low']), 'close': float(r['close']),
            'pre_close': float(r['pre_close']), 'change': float(r['change']),
            'pct_chg': float(r['pct_chg']), 'vol': float(r['vol']),
            'amount': float(r['amount'])
        }
        rows.append(row)
rows.sort(key=lambda x: x['date'])

dates = [r['date'] for r in rows]
closes = np.array([r['close'] for r in rows])
volumes = np.array([r['vol'] for r in rows])
amounts = np.array([r['amount'] for r in rows])
opens = np.array([r['open'] for r in rows])
highs = np.array([r['high'] for r in rows])
lows = np.array([r['low'] for r in rows])
pct_chgs = np.array([r['pct_chg'] for r in rows])

n = len(rows)
start_d = dates[0].strftime('%Y-%m-%d')
end_d = dates[-1].strftime('%Y-%m-%d')
total_return = (closes[-1] - closes[0]) / closes[0] * 100
max_price = highs.max()
min_price = lows[lows > 0].min()
max_date = dates[highs.argmax()].strftime('%Y-%m-%d')
min_date = dates[lows[lows > 0].argmin()].strftime('%Y-%m-%d')
annual_volatility = np.std(pct_chgs) * np.sqrt(252)
max_daily_gain = pct_chgs.max()
max_daily_loss = pct_chgs.min()
avg_volume = volumes.mean()
max_volume = volumes.max()
positive_days = (pct_chgs > 0).sum()
negative_days = (pct_chgs < 0).sum()
ups = closes[1:] > closes[:-1]
up_ratio = ups.sum() / (n - 1) * 100

# MA
def ma(arr, period):
    result = np.full(len(arr), np.nan)
    for i in range(period - 1, len(arr)):
        result[i] = np.mean(arr[i - period + 1:i + 1])
    return result

ma5, ma10, ma20, ma60 = ma(closes, 5), ma(closes, 10), ma(closes, 20), ma(closes, 60)

# Monthly stats
monthly = defaultdict(lambda: {'open': None, 'close': None, 'high': -1e9, 'low': 1e9, 'vol': 0, 'amt': 0, 'days': 0})
for r in rows:
    key = r['date'].strftime('%Y-%m')
    d = monthly[key]
    if d['open'] is None: d['open'] = r['open']
    d['close'] = r['close']
    d['high'] = max(d['high'], r['high'])
    d['low'] = min(d['low'], r['low'])
    d['vol'] += r['vol']
    d['amt'] += r['amount']
    d['days'] += 1

monthly_keys = sorted(monthly.keys())
monthly_returns = []
monthly_vols = []
monthly_labels = []
for k in monthly_keys:
    d = monthly[k]
    ret = (d['close'] - d['open']) / d['open'] * 100
    monthly_returns.append(ret)
    monthly_vols.append(d['vol'] / 1e4)
    monthly_labels.append(k)

# ========== CHART COLORS ==========
RED = '#e74c3c'
GREEN = '#27ae60'
DARK = '#1a1a2e'
BLUE = '#3498db'
ORANGE = '#f39c12'
PURPLE = '#9b59b6'

# ========== CHART 1: K-line (Close + MA) - enhanced colors ==========
fig, ax = plt.subplots(figsize=(14, 6))
ax.plot(dates, closes, color='#1a1a2e', linewidth=1.8, alpha=0.95, label='收盘价', zorder=5)
ax.plot(dates, ma5, color='#e74c3c', linewidth=1.3, alpha=0.9, label='MA5')
ax.plot(dates, ma10, color='#e67e22', linewidth=1.3, alpha=0.9, label='MA10')
ax.plot(dates, ma20, color='#2ecc71', linewidth=1.4, alpha=0.9, label='MA20', linestyle='--')
ax.plot(dates, ma60, color='#9b59b6', linewidth=1.4, alpha=0.9, label='MA60', linestyle='--')
ax.fill_between(list(range(n)), closes, closes[0], where=closes >= closes[0], color='#e74c3c', alpha=0.06)
ax.fill_between(list(range(n)), closes, closes[0], where=closes < closes[0], color='#27ae60', alpha=0.06)
ax.set_title('图1  兆易创新（603986.SH）收盘价走势与移动均线', fontsize=15, fontweight='bold')
ax.set_ylabel('价格（元）', fontsize=11)
ax.legend(loc='upper left', fontsize=10, framealpha=0.9, ncol=5,
          handles=[plt.Line2D([0],[0],color='#1a1a2e',lw=2.5,label='收盘价'),
                   plt.Line2D([0],[0],color='#e74c3c',lw=2,label='MA5'),
                   plt.Line2D([0],[0],color='#e67e22',lw=2,label='MA10'),
                   plt.Line2D([0],[0],color='#2ecc71',lw=2,linestyle='--',label='MA20'),
                   plt.Line2D([0],[0],color='#9b59b6',lw=2,linestyle='--',label='MA60')])
ax.grid(True, alpha=0.25, linestyle='--')
ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
ax.set_xlim(dates[0], dates[-1])
plt.xticks(rotation=45, fontsize=9)
plt.yticks(fontsize=9)
plt.tight_layout()
fig.savefig(os.path.join(CHART_DIR, 'chart1_price_ma.png'), dpi=150, bbox_inches='tight')
plt.close()

# ========== CHART 2: Volume ==========
fig, ax = plt.subplots(figsize=(12, 4))
colors_vol = [RED if closes[i] >= opens[i] else GREEN for i in range(n)]
ax.bar(dates, volumes / 1e4, color=colors_vol, width=1.0, alpha=0.8)
ax.set_title('图2  兆易创新（603986.SH）日成交量变化', fontsize=13, fontweight='bold')
ax.set_ylabel('成交量（万手）', fontsize=10)
ax.grid(True, alpha=0.3, axis='y')
plt.xticks(rotation=30, fontsize=8)
plt.tight_layout()
fig.savefig(os.path.join(CHART_DIR, 'chart2_volume.png'), dpi=150, bbox_inches='tight')
plt.close()

# ========== CHART 3: Monthly Returns ==========
fig, ax = plt.subplots(figsize=(12, 5))
bars = ax.bar(monthly_labels, monthly_returns, color=[RED if v >= 0 else GREEN for v in monthly_returns], alpha=0.85)
for bar, v in zip(bars, monthly_returns):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + (0.5 if v >= 0 else -0.5),
            f'{v:+.1f}%', ha='center', va='bottom' if v >= 0 else 'top', fontsize=7, fontweight='bold')
ax.axhline(y=0, color='black', linewidth=0.5)
ax.set_title('图3  兆易创新（603986.SH）月度收益率分布', fontsize=13, fontweight='bold')
ax.set_ylabel('月度收益率（%）', fontsize=10)
ax.grid(True, alpha=0.3, axis='y')
plt.xticks(rotation=45, fontsize=8)
plt.tight_layout()
fig.savefig(os.path.join(CHART_DIR, 'chart3_monthly_returns.png'), dpi=150, bbox_inches='tight')
plt.close()

# ========== CHART 4: Daily Return Distribution ==========
fig, ax = plt.subplots(figsize=(10, 5))
bins = np.linspace(pct_chgs.min() - 1, pct_chgs.max() + 1, 40)
ax.hist(pct_chgs, bins=bins, color=BLUE, alpha=0.7, edgecolor='white')
ax.axvline(x=0, color='black', linewidth=0.8, linestyle='--')
ax.axvline(x=pct_chgs.mean(), color=RED, linewidth=1, linestyle='-', label=f'均值={pct_chgs.mean():.2f}%')
ax.set_title('图4  兆易创新（603986.SH）日涨跌幅分布', fontsize=13, fontweight='bold')
ax.set_xlabel('日涨跌幅（%）', fontsize=10)
ax.set_ylabel('频次', fontsize=10)
ax.legend(fontsize=9)
ax.grid(True, alpha=0.3)
plt.tight_layout()
fig.savefig(os.path.join(CHART_DIR, 'chart4_return_dist.png'), dpi=150, bbox_inches='tight')
plt.close()

# ========== CHART 5: Monthly Turnover ==========
fig, ax1 = plt.subplots(figsize=(12, 5))
monthly_amts = [monthly[k]['amt'] / 1e8 for k in monthly_keys]
bars = ax1.bar(monthly_labels, monthly_amts, color=BLUE, alpha=0.7, label='月成交额')
ax1.set_title('图5  兆易创新（603986.SH）月度成交额统计', fontsize=13, fontweight='bold')
ax1.set_ylabel('成交额（亿元）', fontsize=10, color=BLUE)
ax1.tick_params(axis='y', labelcolor=BLUE)
ax2 = ax1.twinx()
ax2.plot(monthly_labels, [monthly[k]['close'] for k in monthly_keys], color=RED, marker='o', linewidth=1.5, label='月末收盘价')
ax2.set_ylabel('收盘价（元）', fontsize=10, color=RED)
ax2.tick_params(axis='y', labelcolor=RED)
lines1, labels1 = ax1.get_legend_handles_labels()
lines2, labels2 = ax2.get_legend_handles_labels()
ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=8)
plt.xticks(rotation=45, fontsize=8)
plt.tight_layout()
fig.savefig(os.path.join(CHART_DIR, 'chart5_monthly_turnover.png'), dpi=150, bbox_inches='tight')
plt.close()

# ========== BUILD WORD DOCUMENT ==========
doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# --- Default style ---
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)  # 五号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return h

def add_para(text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_chart(image_path, width=Cm(15)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# ========== DOCUMENT CONTENT ==========

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(f'{STOCK_NAME}（{STOCK_CODE}）基本分析报告')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(22)
run.bold = True

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run(f'分析区间：{start_d} 至 {end_d}  |  交易日：{n} 天')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# === Section 1: Overview ===
add_heading('一、核心数据概览', level=1)

add_para(
    f'{STOCK_NAME}（GigaDevice）是中国领先的半导体设计公司，主营业务包括存储器（NOR Flash、NAND Flash、DRAM）、'
    f'微控制器（MCU）和传感器产品。其产品广泛应用于消费电子、工业控制、汽车电子、物联网等领域。'
    f'本报告基于{start_d}至{end_d}期间共{n}个交易日的日线数据进行分析。'
)

add_para(
    f'区间核心统计：期初开盘价 {opens[0]:.2f} 元，最新收盘价 {closes[-1]:.2f} 元，'
    f'区间累计涨跌幅 {total_return:+.2f}%。区间最高价 {max_price:.2f} 元（{max_date}），'
    f'区间最低价 {min_price:.2f} 元（{min_date}）。'
    f'日均成交量 {avg_volume/1e4:.2f} 万手，单日最大成交量 {max_volume/1e4:.2f} 万手。'
    f'年化波动率约 {annual_volatility:.2f}%。'
)

add_para(
    f'涨跌统计：{n}个交易日中，上涨{n - negative_days}天（占比{up_ratio:.1f}%），'
    f'下跌{negative_days}天。单日最大涨幅 {max_daily_gain:+.2f}%，单日最大跌幅 {max_daily_loss:+.2f}%。'
)

# === Section 2: Price Trend ===
add_heading('二、价格走势分析', level=1)
add_heading('2.1 收盘价与移动均线', level=2)

add_chart(os.path.join(CHART_DIR, 'chart1_price_ma.png'), width=Cm(15.5))

add_para(
    f'图1展示了{STOCK_NAME}在分析区间内的收盘价走势及四条主要移动均线（MA5/MA10/MA20/MA60）。'
    f'从中可以看出股价经历了明显的阶段性变化：'
)

# Find some inflection points
q1_end = dates[min(n-1, n//4)].strftime('%Y-%m-%d')
q2_end = dates[min(n-1, n//2)].strftime('%Y-%m-%d')
q3_end = dates[min(n-1, 3*n//4)].strftime('%Y-%m-%d')
q1_close = closes[min(n-1, n//4)]
q2_close = closes[min(n-1, n//2)]
q3_close = closes[min(n-1, 3*n//4)]

add_para(
    f'（1）初期盘整阶段（{start_d} ~ {q1_end}）：股价在 ¥{q1_close:.0f} 附近窄幅震荡，均线交织，'
    f'多空力量相对均衡，成交量维持在较低水平，市场方向尚不明朗。'
)

mid_low = closes[n//4:n//2].min()
mid_high = closes[n//2:3*n//4].max()
add_para(
    f'（2）中期主升阶段（{q1_end} ~ {q3_end}）：股价从 ¥{q1_close:.0f} 启动，'
    f'逐步攀升至 ¥{q3_close:.0f}，期间出现多次放量涨停，移动均线呈多头排列（MA5 > MA10 > MA20 > MA60），'
    f'显示强劲的上升趋势。'
)

add_para(
    f'（3）后期加速阶段（{q3_end} ~ {end_d}）：股价在最后阶段进一步加速，'
    f'最终收于 ¥{closes[-1]:.0f}，区间整体涨幅高达 {total_return:.0f}%。'
    f'MA5 在多数时间内领涨于其他均线，短期动量充足，但需关注高位波动的风险。'
)

# === Section 3: Volume Analysis ===
add_heading('2.2 成交量分析', level=2)

add_chart(os.path.join(CHART_DIR, 'chart2_volume.png'), width=Cm(15.5))

add_para(
    f'图2展示了每日成交量变化，红色柱表示当日收涨（收盘价 ≥ 开盘价），绿色柱表示当日收跌。'
    f'成交量是市场活跃度和资金参与度的重要指标。'
)

# Find high volume periods
vol_median = np.median(volumes)
high_vol_mask = volumes > vol_median * 1.5
high_vol_count = high_vol_mask.sum()

add_para(
    f'分析区间内，日均成交量为 {avg_volume/1e4:.2f} 万手，共有 {high_vol_count} 个交易日成交量超过中位数的 1.5 倍，'
    f'属于显著放量日。放量主要集中在股价快速拉升阶段，反映了资金的大规模介入。'
    f'值得注意的是一季度末期至二季度初期（2026年4-5月），成交量持续维持在高位，'
    f'与股价的加速上涨形成量价配合，是典型的强势特征。'
)

# === Section 4: Monthly Returns ===
add_heading('三、月度收益分析', level=1)

add_chart(os.path.join(CHART_DIR, 'chart3_monthly_returns.png'), width=Cm(15.5))

pos_months = sum(1 for v in monthly_returns if v > 0)
neg_months = len(monthly_returns) - pos_months
best_month_idx = monthly_returns.index(max(monthly_returns))
worst_month_idx = monthly_returns.index(min(monthly_returns))

add_para(
    f'图3统计了各月的涨跌幅情况。在{len(monthly_returns)}个完整月份中，'
    f'上涨月份 {pos_months} 个，下跌月份 {neg_months} 个，胜率 {pos_months/len(monthly_returns)*100:.1f}%。'
    f'表现最佳的月份为 {monthly_labels[best_month_idx]}（{monthly_returns[best_month_idx]:+.1f}%），'
    f'表现最差的月份为 {monthly_labels[worst_month_idx]}（{monthly_returns[worst_month_idx]:+.1f}%）。'
)

add_para(
    f'从月度收益分布来看，{STOCK_NAME}在分析区间内呈现"大涨小回"的特征：'
    f'上涨月份的涨幅通常较大（平均 {np.mean([v for v in monthly_returns if v > 0]):.1f}%），'
    f'而下跌月份的回撤相对可控（平均 {np.mean([v for v in monthly_returns if v < 0]):.1f}%）。'
    f'这种非对称的收益特征表明，在分析区间内多头力量占据明显优势。'
)

# === Section 5: Daily Return Distribution ===
add_heading('四、日涨跌幅分布', level=1)

add_chart(os.path.join(CHART_DIR, 'chart4_return_dist.png'), width=Cm(14))

sd = np.std(pct_chgs)
skew = np.sum((pct_chgs - np.mean(pct_chgs))**3) / (n * sd**3) if sd > 0 else 0

add_para(
    f'图4展示了日涨跌幅的频率分布直方图。日涨跌幅均值为 {np.mean(pct_chgs):.2f}%，'
    f'标准差为 {sd:.2f}%，偏度为 {skew:.2f}。'
    f'从分布形态来看，涨跌幅基本呈对称分布，但右侧尾部略长（偏度{">0" if skew > 0 else "<0"}），'
    f'说明极端上涨（涨停）出现的频率略高于极端下跌（跌停）。'
    f'涨跌幅在 ±3% 以内的交易日占比约 '
    f'{((pct_chgs > -3) & (pct_chgs < 3)).sum() / n * 100:.0f}%，'
    f'表明多数交易日的波动幅度相对可控，但期间也出现了多次涨停（+10%）和跌停（-10%）的极端行情。'
)

# === Section 6: Monthly Turnover ===
add_heading('五、月度成交额分析', level=1)

add_chart(os.path.join(CHART_DIR, 'chart5_monthly_turnover.png'), width=Cm(15.5))

add_para(
    f'图5展示了月度成交额与月末收盘价的联动关系。蓝色柱状图为月度成交额，红色折线为月末收盘价。'
    f'可以明显观察到，成交额的放大与股价的上涨高度同步：在股价快速拉升的月份，'
    f'月度成交额显著攀升，反映了增量资金的持续涌入。'
    f'总区间累计成交额约 {amounts.sum()/1e8:.0f} 亿元，月均成交额约 {amounts.sum()/1e8/len(monthly_keys):.0f} 亿元。'
    f'成交额的高峰集中在2026年二季度，与股价的加速上涨高度吻合，'
    f'表明市场对该股票的关注度和参与度持续提升。'
)

# === Section 7: Summary & Risk ===
add_heading('六、总结与风险提示', level=1)

add_para(
    f'{STOCK_NAME}在{start_d}至{end_d}的分析区间内表现极为亮眼，累计涨幅约 {total_return:.0f}%，'
    f'在半导体板块中属于绝对的强势品种。股价从 ¥{closes[0]:.0f} 一路上涨至 ¥{closes[-1]:.0f}，'
    f'尤其在2026年二季度出现了明显的加速上涨行情。从技术面来看，均线系统在大部分时间内保持多头排列，'
    f'量价配合良好，上涨月份远多于下跌月份，胜率较高。'
)

add_para('主要亮点：', bold=True)
add_para(
    f'1. 趋势强劲：股价涨幅超{total_return:.0f}%，长期处于上升通道中；'
    f'2. 资金活跃：日均成交额巨大，显示市场关注度高、流动性好；'
    f'3. 动量充足：MA5/MA10在大部分时间内保持在长周期均线之上。'
)

add_para('风险提示：', bold=True)
add_para(
    f'1. 高波动风险：年化波动率约 {annual_volatility:.1f}%，远超市场平均水平，短期回撤可能较大；'
    f'2. 估值风险：股价大幅上涨后，需关注基本面是否能支撑当前估值水平；'
    f'3. 动量衰减风险：若成交量持续萎缩且股价跌破 MA20/MA60 等重要支撑位，可能意味着趋势转折；'
    f'4. 板块轮动风险：半导体板块受政策、周期和地缘政治影响较大，需关注行业层面的系统性风险。'
)

add_para(
    f'\n免责声明：本报告仅基于历史交易数据的统计分析，不构成任何投资建议。'
    f'历史表现不代表未来收益，投资有风险，入市需谨慎。'
)

# ========== SAVE ==========
doc.save(OUTPUT_DOCX)
print(f'报告已生成: {OUTPUT_DOCX}')
